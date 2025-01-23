# 安装插件    pip install python-docx


import random
from docx import Document
from docx.shared import Pt
import datetime

# 生成一个随机的计算式子
def generate_expression():
    num1 = round(random.uniform(1, 100), random.choice([1, 2]))  # 生成1-2位小数的数
    num2 = round(random.uniform(1, 100), random.choice([1, 2]))
    return num1, num2, f"{num1} × {num2}  = "

# 创建一个新的Word文档
doc = Document()

# 创建一个4行3列的表格来展示式子
table = doc.add_table(rows=4, cols=3)

# 用来存储所有答案的列表
answers = []

# 设置表格样式，使表格内容均匀分布
for row in range(4):
    for col in range(3):
        num1, num2, expression = generate_expression()  # 生成一个计算式
        answer = round(num1 * num2, 2)  # 计算答案并保留两位小数
        answers.append(answer)  # 将答案存入数组

        cell = table.cell(row, col)
        
        # 插入计算式
        p1 = cell.add_paragraph(expression)
        p1.alignment = 1  # 计算式居中对齐
        run1 = p1.runs[0]
        run1.font.size = Pt(14)  # 设置字体大小

        # 插入一个空白行让学生填写答案
        p2 = cell.add_paragraph()  # 这里是留给学生填写的空白行
        p2.alignment = 1  # 空白处居中对齐
        run2 = p2.add_run('')  # 确保空白行有一个运行
        run2.font.size = Pt(14)  # 设置字体大小
        p2.add_run("\n\n\n\n")  # 让空白部分下方有更多空白







# 调整表格的列宽，使式子均匀分布
for col in range(3):
    table.columns[col].width = Pt(150)

# 添加答案表格
answers_table = doc.add_table(rows=4, cols=3)
answer_index = 0  # 初始化答案索引
for row in range(4):
    for col in range(3):
        # 从答案数组中获取预先计算的答案
        answer = answers[answer_index]
        answer_index += 1

        # 将答案插入新的表格中
        cell = answers_table.cell(row, col)
        cell.text = str(answer)  # 直接设置单元格的文本内容

        # 设置字体大小和居中对齐
        cell.paragraphs[0].alignment = 1  # 答案居中对齐
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(12)  # 设置字体大小，确保紧凑
        
        # ============带回车的答案
        # cell = answers_table.cell(row, col)
        # p = cell.add_paragraph(str(answer))  # 插入答案
        # p.alignment = 1  # 答案居中对齐
        # run = p.runs[0]
        # run.font.size = Pt(12)  # 设置字体大小，确保紧凑

# 设置答案表格紧凑并放置在页面底部
answers_table.style = 'Table Grid'
answers_table.autofit = True

# 获取当前日期和时间
now = datetime.datetime.now()
# 格式化日期和时间
formatted_time = now.strftime("%Y%m%d%H%M")

# 保存文档
doc.save(f'test{formatted_time}.docx')

print("The Word file has been generated！")
